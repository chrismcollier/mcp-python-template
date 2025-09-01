import os
import json
from typing import Dict, Any, List, Optional
# import pandas as pd
# import numpy as np
import requests
# import plotly.express as px
# import seaborn as sns
# import scipy.stats as stats
# import matplotlib.pyplot as plt
# import scikit-learn as sklearn
# import statsmodels.api as sm
# import torch
import anthropic
import shutil
from docx import Document
from docx.shared import Inches
from datetime import datetime
import PyPDF2
import docx2txt



# Function for reading uploaded training documents
def read_training_documents(upload_folder: str = "uploads") -> str:
    """
    Reads all uploaded training documents and combines them into a single text.
    
    Args:
        upload_folder (str): Path to the folder containing uploaded documents
    
    Returns:
        str: Combined text content from all uploaded documents
    """
    training_content = ""
    
    if not os.path.exists(upload_folder):
        return ""
    
    for filename in os.listdir(upload_folder):
        file_path = os.path.join(upload_folder, filename)
        
        try:
            if filename.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    training_content += f"\n\n--- {filename} ---\n"
                    training_content += f.read()
                    
            elif filename.endswith('.docx'):
                text = docx2txt.process(file_path)
                training_content += f"\n\n--- {filename} ---\n"
                training_content += text
                
            elif filename.endswith('.pdf'):
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text()
                    training_content += f"\n\n--- {filename} ---\n"
                    training_content += text
                    
        except Exception as e:
            print(f"Error reading {filename}: {e}")
            continue
    
    return training_content

def personalize_outline(generic_outline: str, child_info: str, book_title: str, model: str = "claude-sonnet-4-20250514") -> str:
    """
    Personalizes a generic story outline using child-specific information via Claude.

    Args:
        generic_outline (str): The base story structure to customize.
        child_info (str): Details about the child (name, age, traits, interests, etc.).
        model (str): The Claude model to use. Defaults to Claude 3 Sonnet.

    Returns:
        str: A tailored story outline.
    """
    client = anthropic.Anthropic()

    prompt = (
        f"You are a storytelling assistant. Given a generic children's chapter book outline for a book called {book_title} "
        "and information about a child, personalize the story to suit that child. Replace character names, "
        "adjust motivations, settings, and themes to align with the child's personality, interests, and age.\n\n"
        f"### CHILD INFO:\n{child_info}\n\n"
        f"### GENERIC OUTLINE:\n{generic_outline}\n\n"
        "### PERSONALIZED OUTLINE:"
    )

    response = client.messages.create(
        model=model,
        max_tokens=1024,
        messages=[
            {"role": "user", "content": prompt}
        ]
    )

    return response.content[0].text.strip()


# Function for generating a chapter with training data integration
def generate_chapter(outline: str, chapter_number: int, model: str = "claude-sonnet-4-20250514") -> str:
    """
    Generates a chapter of a children's chapter book using Claude with training data.

    Args:
        outline (str): The outline of the chapter to generate.
        chapter_number (int): The chapter number to generate.
        model (str): The Claude model to use. Defaults to Claude 3 Sonnet.

    Returns:
        str: A chapter of a children's chapter book.
    """

    client = anthropic.Anthropic()
    
    # Get training content from uploaded documents
    training_content = read_training_documents()
    
    training_section = ""
    if training_content:
        training_section = f"""
        
### TRAINING DOCUMENTS (Use for style, tone, and writing approach):
{training_content[:3000]}  # Limit to avoid token limits
...

Use the above training documents to inform your writing style, tone, and approach.
"""

    prompt = (
        "You are a whimsical and wise children's book author "
        "who crafts imaginative stories with heartfelt lessons "
        "and a dash of silliness. You can write stories, fairy "
        "tales, and full length chapter books. You celebrate "
        "kindness, imagination, courage, and being different.\n\n"
        
        f"{training_section}"
        
        "Follow these guidelines:\n"
        "- Use clear, age-appropriate language\n"
        "- Include dialogue and action\n"
        "- Create engaging descriptions\n"
        "- Build suspense and excitement\n"
        "- Include positive lessons naturally\n"
        "- Make each chapter end with anticipation for the next\n\n"

        "### STORY OUTLINE:\n"
        f"{outline}\n\n"
        "### CHAPTER TO WRITE:\n"
        f"Chapter {chapter_number}\n\n"
        "Write the chapter to be 700-800 words long. Include chapter transitions and engaging storytelling."
    )

    response = client.messages.create(
        model=model,
        max_tokens=2000,  # Increased for longer chapters
        messages=[
            {"role": "user", "content": prompt}
        ]
    )

    return response.content[0].text.strip()

def duplicate_template(template_filename: str, book_title: str) -> str:
    """
    Creates a new book file from the template with the book title as the filename.
    
    Args:
        template_filename (str): Name of the template file (e.g., "template.docx")
        book_title (str): The title of the book (used for filename)
        output_directory (str): Directory to save the new file (optional, defaults to current directory)
    
    Returns:
        str: Path to the newly created book file
    """
    template_filename = "C:\\Users\\Chris Collier\\repos\\chapter-book-generator\\template.docx"
    doc = Document(template_filename)
    print("Template loaded successfully")

    # Clean title for filename
    invalid_chars = '<>:|"?*\\/'
    clean_title = "".join(c if c not in invalid_chars else '_' for c in book_title)
    filename = clean_title.replace(' ', '_')
    print(f"cleaned filename: {filename}")

    # Create output filename
    output_filename = f"C:\\Users\\Chris Collier\\repos\\chapter-book-generator\\{filename}.docx"
    print(f"output filename: {output_filename}")

    # Copy template to new file
    doc.save(output_filename)
    print("Template saved successfully")

    return output_filename

def inject_chapter(book_file_name: str, chapter_content: str, chapter_number: int, chapter_title: str) -> str:
    """
    Injects a chapter into the book file created by the duplicate template function.

    Args:
        book_file_name (str): the name of the book file created through the duplicate template tool
        chapter_content (str): the content of the chapter to be injected
        chapter_number (int): the number of the chapter to be injected
        chapter_title (str): the title of the chapter to be injected
    
    Returns:
        str: the content of the book file with the chapter injected along with confirmation that the document has been saved
    """
    doc = Document(book_file_name)
    print("Book file loaded successfully")

    # define placeholders to search for
    title_placeholder = f"Chapter {chapter_number} name"
    content_placeholder = f"Chapter {chapter_number} text"
    title_check = f"Chapter {chapter_number}:"

    title_replaced = False
    content_replaced = False
    
    chapter_title = chapter_title.replace(': ','')

    # Stop searching once both placeholders are found and replaced
    for paragraph in doc.paragraphs:
        # Handle title replacement while preserving formatting
        if not title_replaced and title_placeholder in paragraph.text:
            # Store original formatting
            original_style = paragraph.style
            original_alignment = paragraph.alignment
            original_left_indent =  paragraph.paragraph_format.left_indent
            original_right_indent = paragraph.paragraph_format.left_indent
            
            # Get font formatting from first run if it exists
            original_font_name = None
            original_font_size = None
            original_bold = None
            if paragraph.runs:
                original_font_name = paragraph.runs[0].font.name
                original_font_size = paragraph.runs[0].font.size
                original_bold = paragraph.runs[0].bold

            # Clear and rebuild paragraph
            paragraph.clear()
            run = paragraph.add_run(chapter_title)
            
            # Reapply formatting
            paragraph.style = original_style
            paragraph.alignment = original_alignment
            if original_font_name:
                run.font.name = original_font_name
            if original_font_size:
                run.font.size = original_font_size
            if original_bold is not None:
                run.bold = original_bold
            if original_left_indent:
                run.paragraph_format.left_indent = original_left_indent
            if original_right_indent:
                run.paragraph_format.right_indent = original_right_indent
                
            title_replaced = True
        
        # Handle content replacement with paragraph formatting
        elif not content_replaced and content_placeholder in paragraph.text and title_check not in paragraph.text:
            # Store original formatting
            original_style = paragraph.style
            original_font_name = None
            original_font_size = None
            if paragraph.runs:
                original_font_name = paragraph.runs[0].font.name
                original_font_size = paragraph.runs[0].font.size
            

            # Clear and rebuild paragraph
            paragraph.clear()

            content_paragraphs = chapter_content.split('\n\n')

            # Handle first paragraph in current location
            first_para_run = paragraph.add_run(content_paragraphs[1])
            paragraph.style = original_style
            paragraph.paragraph_format.first_line_indent = Inches(0.5)

            # Apply font formatting
            if original_font_name:
                first_para_run.font.name = original_font_name
            if original_font_size:
                first_para_run.font.size = original_font_size

            # CAPTURE ALL PARAGRAPH FORMATTING from the original
            original_para_format = paragraph.paragraph_format

            # Add additional paragraphs right after the current paragraph
            last_inserted = paragraph
            for content_para in content_paragraphs[2:]:
                # Create new paragraph
                new_para = doc.add_paragraph()
                new_run = new_para.add_run(content_para)
                new_para.style = original_style
                
                # COPY ALL FORMATTING from original paragraph
                new_para.paragraph_format.first_line_indent = original_para_format.first_line_indent
                new_para.paragraph_format.left_indent = original_para_format.left_indent
                new_para.paragraph_format.right_indent = original_para_format.right_indent
                
                # Apply font formatting
                if original_font_name:
                    new_run.font.name = original_font_name
                if original_font_size:
                    new_run.font.size = original_font_size
                
                # Move the new paragraph to the correct position
                last_inserted._element.addnext(new_para._element)
                last_inserted = new_para
            content_replaced = True
        
        # Exit early if both replacements are done
        if title_replaced and content_replaced:
            break
    
    doc.save(book_file_name)
    return f"Successfully injected Chapter {chapter_number}: '{chapter_title}'"


def extract_chapter_title(chapter_content: str, chapter_number: int) -> str:
    """
    Extracts a chapter title from the generated chapter content.
    
    Args:
        chapter_content (str): The generated chapter content
        chapter_number (int): The chapter number
    
    Returns:
        str: Extracted or generated chapter title
    """
    lines = chapter_content.split('\n')
    
    # Look for common title patterns in first few lines
    for line in lines[:5]:
        line = line.strip()
        if line and (
            line.startswith('Chapter') or 
            line.startswith('#') or
            (len(line) < 50 and len(line.split()) <= 6)
        ):
            # Clean up the title
            title = line.replace('#', '').replace('Chapter ' + str(chapter_number), '').strip()
            if title and len(title) > 2:
                return title
    
    # If no good title found, generate a simple one
    return f"Chapter {chapter_number}"

def generate_book(generic_outline: str, child_info: str, book_title: str, model: str = "claude-sonnet-4-20250514") -> str:
    """
    Generates a complete 10-chapter children's book and saves it to a document.

    Args:
        generic_outline (str): The generic story outline to personalize.
        child_info (str): Information about the child for personalization.
        book_title (str): The title of the book.
        model (str): The Claude model to use. Defaults to Claude 3 Sonnet.

    Returns:
        str: The complete book preview and confirmation of file creation.
    """
    
    # personalize the outline based on the child's info and the book title
    personalized_outline = personalize_outline(generic_outline, child_info, book_title)
    
    # create the book file
    book_file_name = duplicate_template("template.docx", book_title)
    
    print(f"Starting generation of '{book_title}' with 10 chapters...")
    
    # Initialize the book content
    book_content = []
    book_content.append(f"# {book_title}")
    book_content.append(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
    book_content.append("=" * 50)
    book_content.append("")
    
    # Generate each chapter (now 1-10 instead of 1-9)
    for chapter_num in range(1, 11):
        print(f"Generating Chapter {chapter_num}...")
        
        try:
            chapter_content = generate_chapter(
                outline=personalized_outline, 
                chapter_number=chapter_num, 
                model=model
            )
            
            # Extract chapter title from content
            chapter_title = extract_chapter_title(chapter_content, chapter_num)
            
            # Add chapter to book
            book_content.append(f"## Chapter {chapter_num}: {chapter_title}")
            book_content.append("")
            book_content.append(chapter_content)
            book_content.append("")
            book_content.append("=" * 30)
            book_content.append("")

            # inject the chapter into the book file
            inject_chapter(book_file_name, chapter_content, chapter_num, chapter_title)

            print(f"✓ Chapter {chapter_num}: '{chapter_title}' completed ({len(chapter_content)} characters)")
            
        except Exception as e:
            error_msg = f"Error generating Chapter {chapter_num}: {str(e)}"
            print(f"✗ {error_msg}")
            book_content.append(f"## Chapter {chapter_num}")
            book_content.append("")
            book_content.append(f"[{error_msg}]")
            book_content.append("")
            book_content.append("=" * 30)
            book_content.append("")
    
    # Combine all content
    full_book_text = "\n".join(book_content)
    
    return f"✅ Successfully generated '{book_title}' with 10 chapters!\n\nBook preview:\n{full_book_text[:500]}..."